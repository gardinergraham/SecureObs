from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

from build_nurse_manual import (  # noqa: E402
    ASSETS,
    LIGHT_BLUE,
    LIGHT_TEAL,
    MUTED,
    NAVY,
    OUT as _NURSE_OUT,
    TEAL,
    WARN,
    add_bullets,
    add_callout,
    add_header_footer,
    add_screenshot,
    add_step_blocks,
    add_two_col_table,
    set_cell_border,
    set_cell_margins,
    set_cell_shading,
    set_run,
    set_table_width,
    style_doc,
)


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "website" / "training" / "SecureObs_HCF_Security_Training_Manual.docx"


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    table.columns[0].width = Inches(3.9)
    table.columns[1].width = Inches(2.9)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, "FFFFFF", "0")
            set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(10)
    p.add_run().add_picture(str(ASSETS / "secureobs-logo-new.png"), width=Inches(1.05))

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("HCF AND SECURITY STAFF GUIDE")
    set_run(r, size=10, color=TEAL, bold=True)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 0.95
    r = p.add_run("Safe ward support and security checks")
    set_run(r, size=30, color=NAVY, bold=True)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run(
        "A practical SecureObs guide for healthcare support staff and security staff using observations, enhanced support, security checks, safety escalation and tasks."
    )
    set_run(r, size=11, color=MUTED, bold=True)

    chip_table = left.add_table(rows=0, cols=1)
    set_table_width(chip_table, 5200)
    for chip in [
        "Sign in and confirm ward context",
        "Support observations and enhanced observation cover",
        "Record security checks and count variances",
        "Raise safety incidents and patient tasks",
        "Escalate clearly to nurse in charge",
    ]:
        row = chip_table.add_row()
        cell = row.cells[0]
        set_cell_shading(cell, LIGHT_TEAL)
        set_cell_border(cell, "BDECF2")
        set_cell_margins(cell, top=80, bottom=80, start=150, end=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(chip)
        set_run(r, size=9.5, color=NAVY, bold=True)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS / "screenshots2" / "security_Checks_Screen.jpg"), width=Inches(2.75))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(8)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Version 1.1 • September 2026 • For HCF, support workers, security staff and ward trainers")
    set_run(r, size=9, color=MUTED, bold=True)


def build() -> None:
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)
    doc.sections[0].header.paragraphs[0].clear()
    set_run(doc.sections[0].header.paragraphs[0].add_run("SecureObs HCF and Security Staff Training Manual"), size=8.5, color=MUTED, bold=True)
    doc.sections[0].footer.paragraphs[0].clear()
    doc.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(doc.sections[0].footer.paragraphs[0].add_run("SecureObs • HCF, security checks and ward support"), size=8, color=MUTED, bold=True)
    add_cover(doc)
    doc.add_page_break()

    doc.add_heading("1. What this route covers", level=1)
    add_callout(
        doc,
        "Role purpose",
        "This route supports safe ward operations. HCF and security staff may help record observations, security checks, incidents and tasks, but clinical escalation remains with the nurse in charge or responsible clinician.",
    )
    add_two_col_table(
        doc,
        [
            ("HCF focus", "Observation support, patient location/presentation updates, enhanced observation support, tasks and escalation."),
            ("Security focus", "Security checks, count variances, searches/checkpoints where configured, safety incidents and escalation."),
            ("Always escalate", "Any immediate risk, uncertainty, safeguarding issue, medication concern or clinical deterioration must be escalated to nursing/clinical staff."),
        ],
        header=("Area", "Guidance"),
    )

    doc.add_heading("2. Start of shift", level=1)
    add_step_blocks(
        doc,
        [
            ("Sign in", "Use NFC staff card or STAFFCODE/PIN. Temporary staff should check the allocation start and end time shown in the current session. If access fails, ask the nurse in charge or manager to check the record."),
            ("Confirm context", "Check the selected site, ward and staff name before recording anything."),
            ("Review ward overview", "Look for overdue checks, active incidents, patient tasks and sync issues."),
            ("Know your allocation", "Confirm which patients, observation levels, security checks or tasks you are responsible for."),
            ("Escalate uncertainty", "If you are unsure whether you should record something, speak to the nurse in charge first."),
        ],
    )

    doc.add_heading("3. Observation support", level=1)
    add_two_col_table(
        doc,
        [
            ("Patient identity", "Check name, room and hospital number before saving any record."),
            ("Location", "Record where the patient is actually seen using the location labels configured for that ward."),
            ("Presentation", "Record awake or asleep accurately."),
            ("Notes", "Keep notes brief and factual. Escalate clinical concerns instead of writing long clinical interpretation."),
            ("Enhanced observations", "If assigned, record location/presentation and follow the plan of care and observation level exactly."),
            ("Late checks", "The record becomes late at the due time. During the first five minutes a reason is not mandatory; after five minutes, record the real reason and inform the nurse in charge if checks are repeatedly late or unsafe."),
            ("NFC or QR verification", "Where enabled, scan the correct room or personal tag and confirm direct visual observation. A room tag does not prove the patient was present."),
        ],
        header=("Action", "Support staff guidance"),
    )
    add_screenshot(doc, ASSETS / "screenshots2" / "general_Obs_Screen-2.jpg", "General observations: select the correct patient, location and presentation.", width=4.75)

    doc.add_heading("4. Security checks", level=1)
    add_step_blocks(
        doc,
        [
            ("Open Security checks", "Use the security checks screen when the ward has this module enabled and you are assigned to complete checks."),
            ("Choose checkpoint", "Select the correct area or patient-specific check from the checkpoint list."),
            ("Enter counts accurately", "For cutlery or item-based checks, enter actual counts for every required item."),
            ("Record variances", "If counts do not match expected values, save the variance and escalate according to local process."),
            ("Use history", "Review history to check what has already been recorded for the day or shift."),
            ("Never hide uncertainty", "If a count is unclear, record the facts and escalate rather than guessing."),
        ],
    )
    add_callout(
        doc,
        "Count variance rule",
        "A count variance should be treated as a safety issue. Record what was counted, preserve the audit trail and tell the nurse in charge or security lead.",
        fill=WARN,
    )
    add_screenshot(doc, ASSETS / "screenshots2" / "Security_Checks.jpg", "Security checks: checkpoint status, recording and history.", width=4.7)

    doc.add_heading("5. Safety incidents and patient tasks", level=1)
    add_two_col_table(
        doc,
        [
            ("Safety incident", "Use for injury, fall, violence/aggression, self-harm, safeguarding, security or other significant concern."),
            ("Severity", "Use RAG rating honestly: red for immediate/serious risk, amber for review, green for monitor."),
            ("Body map", "Select affected areas for injury or physical concern where relevant."),
            ("Immediate action", "Record first aid, area made safe, staff informed, nurse called or other immediate action."),
            ("Patient account", "Record the patient’s own words where possible, or why this could not be obtained."),
            ("Patient task", "Use tasks for follow-up actions, but urgent risks must still be escalated verbally."),
        ],
        header=("Area", "How to record"),
    )

    doc.add_heading("6. What not to do", level=1)
    add_bullets(
        doc,
        [
            "Do not use another staff member’s login.",
            "Do not record an observation if you did not complete or witness it.",
            "Do not guess counts, locations or patient presentation.",
            "Do not resolve incidents or tasks unless you are authorised and the action has genuinely been completed.",
            "Do not use SecureObs as a replacement for emergency escalation, safeguarding policy or nurse-in-charge review.",
        ]
    )

    doc.add_heading("7. End of shift and handover awareness", level=1)
    add_two_col_table(
        doc,
        [
            ("Before leaving", "Check that assigned observations, security checks and tasks are completed or handed over."),
            ("Sync badge", "If pending items remain, tell the nurse in charge and follow local sync process."),
            ("Temporary access warning", "At the five-minute expiry warning, save unfinished work and prepare to hand over. The session ends at the allocation end time."),
            ("Handover notes", "Give factual verbal handover for anything that needs immediate awareness."),
            ("Active incidents", "Do not leave active safety concerns without clear ownership."),
        ],
        header=("End-of-shift item", "Expected action"),
    )

    doc.add_heading("8. Troubleshooting", level=1)
    add_two_col_table(
        doc,
        [
            ("Cannot sign in", "Check STAFFCODE/PIN or NFC tag. Ask nurse in charge/manager to confirm active status and ward access."),
            ("Security area missing", "The ward may not have the check configured or enabled. Escalate to manager/security lead."),
            ("Patient missing", "Confirm selected ward and ask nursing staff to check patient setup."),
            ("Sync issue appears", "If trained, open sync status, read the waiting item and retry after authorised sign-in. Otherwise tell the nurse in charge."),
            ("Unsure if clinical", "Stop and escalate. Do not record clinical interpretation beyond your role."),
        ],
        header=("Issue", "What to do"),
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
