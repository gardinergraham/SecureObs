from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

from build_nurse_manual import (  # noqa: E402
    ASSETS,
    LIGHT_TEAL,
    MUTED,
    NAVY,
    TEAL,
    WARN,
    add_bullets,
    add_callout,
    add_header_footer,
    add_screenshot,
    add_step_blocks,
    add_two_col_table,
    set_cell_border,
    set_cell_margins,
    set_cell_shading,
    set_run,
    set_table_width,
    style_doc,
)


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "website" / "training" / "SecureObs_Doctor_Prescriber_Training_Manual.docx"


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    table.columns[0].width = Inches(3.9)
    table.columns[1].width = Inches(2.9)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, "FFFFFF", "0")
            set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(10)
    p.add_run().add_picture(str(ASSETS / "secureobs-logo-new.png"), width=Inches(1.05))

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("DOCTOR AND PRESCRIBER GUIDE")
    set_run(r, size=10, color=TEAL, bold=True)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 0.95
    r = p.add_run("Clinical review, prescribing and care-plan oversight")
    set_run(r, size=28, color=NAVY, bold=True)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run(
        "A focused SecureObs guide for doctors and authorised prescribers reviewing patients, prescribing medication, checking clinical evidence and contributing to care plans."
    )
    set_run(r, size=11, color=MUTED, bold=True)

    chip_table = left.add_table(rows=0, cols=1)
    set_table_width(chip_table, 5200)
    for chip in [
        "Sign in and select ward/patient",
        "Review patient dashboard, NEWS2 and history",
        "Prescribe regular, PRN, depot or rapid medication",
        "Record allergies, ADRs and stopped medication",
        "Review notes, care plans and safety incidents",
    ]:
        row = chip_table.add_row()
        cell = row.cells[0]
        set_cell_shading(cell, LIGHT_TEAL)
        set_cell_border(cell, "BDECF2")
        set_cell_margins(cell, top=80, bottom=80, start=150, end=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(chip)
        set_run(r, size=9.5, color=NAVY, bold=True)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS / "screenshots2" / "Meds_Prescribe.jpg"), width=Inches(2.75))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(8)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Version 1.0 • July 2026 • For doctors, prescribers and clinical reviewers")
    set_run(r, size=9, color=MUTED, bold=True)


def build() -> None:
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)
    doc.sections[0].header.paragraphs[0].clear()
    set_run(doc.sections[0].header.paragraphs[0].add_run("SecureObs Doctor and Prescriber Training Manual"), size=8.5, color=MUTED, bold=True)
    doc.sections[0].footer.paragraphs[0].clear()
    doc.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(doc.sections[0].footer.paragraphs[0].add_run("SecureObs • Doctor review, prescribing and care planning"), size=8, color=MUTED, bold=True)
    add_cover(doc)
    doc.add_page_break()

    doc.add_heading("1. What this doctor route covers", level=1)
    add_callout(
        doc,
        "Doctor route purpose",
        "The doctor route is for reviewing clinical evidence, adding medical context, prescribing or stopping medication where authorised, and contributing to care-plan and safety review.",
    )
    add_two_col_table(
        doc,
        [
            ("Who should use this guide", "Doctors, authorised prescribers and clinical reviewers using SecureObs."),
            ("Prescribing access", "Doctor role can prescribe when the staff record allows prescribing. Do not prescribe under another user’s session."),
            ("Clinical responsibility", "SecureObs supports clinical review but does not replace professional judgement, medicines policy or urgent escalation."),
        ],
        header=("Item", "Guidance"),
    )

    doc.add_heading("2. Sign in and review patient context", level=1)
    add_step_blocks(
        doc,
        [
            ("Sign in", "Use NFC staff card or STAFFCODE/PIN and confirm the correct site and ward."),
            ("Select patient", "Choose the patient and confirm room and hospital number before reviewing or prescribing."),
            ("Review dashboard/timeline", "Look at latest observation location, presentation, NEWS2, food/fluid, incidents, tasks and progress trends."),
            ("Read notes and care plans", "Review recent notes and active care-plan goals before adding decisions or changing medication."),
            ("Check safety information", "Review active safety incidents or escalation items relevant to medical review."),
        ],
    )

    doc.add_page_break()
    doc.add_heading("3. Medication chart and prescribing", level=1)
    add_two_col_table(
        doc,
        [
            ("Allergies / ADRs", "Review and update allergies and adverse drug reactions before prescribing."),
            ("Drug approved name", "Enter the medicine name clearly using local policy and approved naming conventions."),
            ("Dose and route", "Record dose and route such as Oral, IM, Depot, S/L or Topical."),
            ("Prescription type", "Choose Regular, PRN, Depot or Rapid tranquillisation as appropriate."),
            ("Administration times", "Select scheduled times for regular medication. PRN/rapid items do not use routine scheduled times."),
            ("Additional instructions", "Add pharmacist advice, monitoring instructions or safe administration context where needed."),
            ("Confirmation", "Confirm the prescription or administration action before it becomes part of the record."),
            ("Stopped medication", "Enter stop date/time and reason when discontinuing medication."),
        ],
        header=("Medication item", "Doctor / prescriber guidance"),
    )
    add_callout(
        doc,
        "Prescribing safety",
        "Always confirm the selected patient before prescribing. Check allergies/ADRs, active prescriptions, recent administrations and any relevant clinical risk before saving.",
        fill=WARN,
    )
    add_screenshot(doc, ASSETS / "screenshots2" / "Medication_Prescribing.jpg", "Medication chart: prescribing, allergies/ADRs and administration grid.", width=4.75)

    doc.add_heading("4. Reviewing medication administration", level=1)
    add_bullets(
        doc,
        [
            "Review given, refused, omitted and unavailable doses in the chart.",
            "Check omission codes and notes when a dose was not given.",
            "Review PRN or rapid tranquillisation records in the context of behaviour, incidents and care plan.",
            "Use medication history to understand recent administration before changing a prescription.",
            "Escalate any discrepancy to the nurse in charge or medication nurse.",
        ]
    )

    doc.add_heading("5. Care-plan and clinical record review", level=1)
    add_two_col_table(
        doc,
        [
            ("Patient notes", "Read recent notes before adding a medical review note. Keep entries factual, attributable and clinically relevant."),
            ("Care plans", "Review identified needs, risks, goals, interventions and patient preferences. Add or request updates where the plan no longer matches clinical need."),
            ("Review date", "Check whether the care plan review date is due or clinically indicated sooner."),
            ("Patient voice", "Consider what matters to the patient, patient preferences and any shared-care/family contribution that has been reviewed by staff."),
            ("Safety incidents", "Review relevant incidents, body map details, immediate action and resolution notes."),
            ("NEWS2 / physical health", "Use NEWS2 and observation history to support clinical assessment and escalation decisions."),
        ],
        header=("Record area", "Doctor review focus"),
    )
    add_screenshot(doc, ASSETS / "screenshots2" / "Patient_Care_Plans_Full.jpg", "Care plans: structured needs, risks, goals, interventions and review date.", width=4.75)

    doc.add_heading("6. Shift handover and team communication", level=1)
    add_two_col_table(
        doc,
        [
            ("Handover awareness", "Review patient summaries where medical context affects the incoming shift."),
            ("Medical action", "Make clear what needs nursing follow-up, observation, physical health review or medication monitoring."),
            ("Tasks", "Create or review patient tasks where follow-up needs ownership."),
            ("Escalation", "Verbally escalate urgent decisions. Do not rely solely on a written note for immediate risk."),
        ],
        header=("Area", "Expected action"),
    )

    doc.add_heading("7. What not to do", level=1)
    add_bullets(
        doc,
        [
            "Do not prescribe or discontinue medication under another staff member’s login.",
            "Do not prescribe without checking allergies/ADRs and active medication records.",
            "Do not use SecureObs as a substitute for urgent clinical escalation.",
            "Do not document vague decisions such as “review later” without ownership or timeframe.",
            "Do not approve family/shared-care information unless it is appropriate and patient-consented according to policy.",
        ]
    )

    doc.add_heading("8. Troubleshooting", level=1)
    add_two_col_table(
        doc,
        [
            ("Cannot prescribe", "Check doctor role, Can prescribe permission and medication chart module status."),
            ("Patient not visible", "Confirm ward selection and ask manager/admin to review patient setup."),
            ("Medication action not saving", "Check sync status and retry; alert nurse in charge if time-critical."),
            ("Wrong patient selected", "Stop before saving. Return to the patient list and confirm identity."),
            ("Care-plan review unclear", "Check recent notes, patient voice, incidents and handover before documenting review outcome."),
        ],
        header=("Issue", "What to do"),
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
