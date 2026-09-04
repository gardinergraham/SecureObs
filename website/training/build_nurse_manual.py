from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "website" / "training" / "SecureObs_Nurse_Route_Training_Manual.docx"
ASSETS = ROOT / "website" / "assets"

NAVY = RGBColor(8, 42, 88)
BLUE = RGBColor(4, 93, 185)
TEAL = RGBColor(8, 167, 181)
MUTED = RGBColor(72, 98, 116)
WHITE = RGBColor(255, 255, 255)
LIGHT_BLUE = "EAF6FB"
LIGHT_TEAL = "E8F8FA"
LIGHT_GREY = "F3F8FA"
WARN = "FFF4D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D6E5EC", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_run(run, *, size=None, color=None, bold=None, italic=None, font="Arial") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(18, 36, 46)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 16, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11, NAVY, 8, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document) -> None:
    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("SecureObs Nurse Route Training Manual")
    set_run(r, size=8.5, color=MUTED, bold=True)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("SecureObs • Nurse route, ward observations and clinical handover")
    set_run(r, size=8, color=MUTED, bold=True)


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    table.columns[0].width = Inches(3.9)
    table.columns[1].width = Inches(2.9)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, "FFFFFF", "0")
            set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(10)
    p.add_run().add_picture(str(ASSETS / "secureobs-logo-new.png"), width=Inches(1.05))

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("NURSE ROUTE TRAINING GUIDE")
    set_run(r, size=10, color=TEAL, bold=True)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 0.95
    r = p.add_run("Working a safe SecureObs shift")
    set_run(r, size=31, color=NAVY, bold=True)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run(
        "A practical nurse workflow manual for recording observations, managing overdue checks, using clinical modules and completing evidence-based handover."
    )
    set_run(r, size=11, color=MUTED, bold=True)

    chip_table = left.add_table(rows=0, cols=1)
    set_table_width(chip_table, 5200)
    for chip in [
        "Start-of-shift dashboard checks",
        "General observations and missed reasons",
        "NEWS2, food/fluid and enhanced observations",
        "Patient notes, care plans and tasks",
        "Safety escalation and shift handover",
    ]:
        row = chip_table.add_row()
        cell = row.cells[0]
        set_cell_shading(cell, LIGHT_TEAL)
        set_cell_border(cell, "BDECF2")
        set_cell_margins(cell, top=80, bottom=80, start=150, end=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(chip)
        set_run(r, size=9.5, color=NAVY, bold=True)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS / "screenshots2" / "General_OBS.jpg"), width=Inches(2.75))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(8)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Version 1.1 • September 2026 • For nurses, nurse in charge, ward staff and clinical trainers")
    set_run(r, size=9, color=MUTED, bold=True)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "C9E5EF")
    set_cell_margins(cell, top=150, bottom=150, start=170, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run(r, size=10.5, color=NAVY, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run(r, size=9.5, color=RGBColor(38, 72, 89))


def add_two_col_table(doc: Document, rows: list[tuple[str, str]], widths=(2.25, 4.2), header=None) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    for col, width in zip(table.columns, widths):
        col.width = Inches(width)
    if header:
        row = table.add_row()
        for i, text in enumerate(header):
            cell = row.cells[i]
            set_cell_shading(cell, "0A2F63")
            set_cell_border(cell, "0A2F63")
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run(r, size=9, color=WHITE, bold=True)
    for label, detail in rows:
        row = table.add_row()
        for i, text in enumerate((label, detail)):
            cell = row.cells[i]
            set_cell_shading(cell, "FFFFFF" if i else LIGHT_GREY)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run(r, size=9, color=NAVY if i == 0 else RGBColor(31, 54, 66), bold=i == 0)


def add_step_blocks(doc: Document, steps: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    table.columns[0].width = Inches(0.65)
    table.columns[1].width = Inches(5.8)
    for idx, (title, body) in enumerate(steps, 1):
        row = table.add_row()
        number_cell, body_cell = row.cells
        set_cell_shading(number_cell, "0A6C85")
        set_cell_border(number_cell, "0A6C85")
        set_cell_margins(number_cell, top=140, bottom=140, start=100, end=100)
        p = number_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(idx))
        set_run(r, size=14, color=WHITE, bold=True)

        set_cell_shading(body_cell, "FFFFFF")
        set_cell_border(body_cell)
        set_cell_margins(body_cell, top=120, bottom=120, start=150, end=150)
        p = body_cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        set_run(r, size=10, color=NAVY, bold=True)
        p = body_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(body)
        set_run(r, size=9, color=RGBColor(45, 70, 82))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run(r, size=9.7, color=RGBColor(30, 55, 68))


def add_screenshot(doc: Document, path: Path, caption: str, width=5.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run(r, size=8.5, color=MUTED, italic=True)


def build() -> None:
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)
    add_cover(doc)
    doc.add_page_break()

    doc.add_heading("1. What the nurse route covers", level=1)
    add_callout(
        doc,
        "Nurse route purpose",
        "The nurse route is designed around a live ward shift: seeing what is due, recording patient checks, adding clinical context, escalating safety concerns and handing over clearly.",
    )
    add_two_col_table(
        doc,
        [
            ("Who should use this guide", "Nurses, nurse in charge, clinical support staff and trainers introducing SecureObs on the ward."),
            ("Before starting", "Confirm you are on the correct site and ward, and that the patient list matches the ward board or local process."),
            ("Clinical record rule", "Record what you observe at the time. Use notes and handover for clinical context, not to replace required observations."),
        ],
        header=("Item", "Guidance"),
    )

    doc.add_heading("2. Start of shift", level=1)
    add_step_blocks(
        doc,
        [
            ("Sign in", "Use NFC staff card or STAFFCODE/PIN. If prompted to change PIN, complete this before starting clinical work. Temporary staff should also check the allocation start and end time shown in the current session."),
            ("Check current session", "Confirm your name, site and ward are correct. Change staff or ward if the wrong context is showing."),
            ("Review ward overview", "Look at due/overdue observations, NEWS2 concerns, incidents, patient tasks, security checks and sync issues."),
            ("Check role allocation", "If rota is enabled, confirm nurse in charge, medication nurse and enhanced observation cover are correct."),
            ("Open patient checks", "Move into General observations when ready to begin recording patient checks."),
        ],
    )
    add_screenshot(doc, ASSETS / "screenshots2" / "General_OBS.jpg", "General observations: patient list, due times and selected patient check panel.", width=4.7)

    doc.add_heading("3. Recording general observations", level=1)
    add_two_col_table(
        doc,
        [
            ("Select patient", "Choose the patient from the list. Check room number and hospital number before saving."),
            ("Current location", "Record where the patient was actually seen using the location labels configured for that ward."),
            ("Presentation", "Record whether the patient was awake or asleep."),
            ("Notes", "Add brief relevant context. Longer clinical narrative should go into Patient notes."),
            ("Save check", "Tap Save check. Confirm the patient row updates and the due time resets."),
            ("If the check is late", "A late check is recorded from the due time. During the first five minutes, save the completed check without a mandatory reason; after five minutes, choose the missed/overdue reason and add context."),
            ("NFC or QR verification", "Where enabled, scan the room or personal tag and confirm that you directly observed the selected patient. If the scan identifies another patient, stop and correct the selection."),
        ],
        header=("Action", "Nurse guidance"),
    )
    add_callout(
        doc,
        "Observation timing",
        "The due time is based on the ward observation interval or the patient's enhanced observation plan. A late record remains visible immediately; the five-minute grace period only delays the mandatory reason. Record the reason honestly once required and escalate according to local policy.",
        fill=WARN,
    )

    doc.add_heading("4. Clinical modules during the shift", level=1)
    add_two_col_table(
        doc,
        [
            ("NEWS2", "Record observations and review score history when NEWS2 is enabled for the ward."),
            ("Food and fluid", "Record meal, snack and drink intake. Low intake or refusal appears in summaries and handover."),
            ("Enhanced observations / TESO", "Use for patients needing higher support, enhanced observation plans or TESO governance."),
            ("Medication chart", "Record administration where your role allows. Review due prompts, allergies/ADRs and medication history."),
            ("Patient notes", "Add longer clinical notes, filter by staff/date and export selected notes as PDF when authorised."),
            ("Care plans", "Create or review structured care plans with needs, risks, goals, interventions, patient preferences and review date."),
            ("Patient dashboard", "Review patient progress and visual summaries across clinical activity."),
            ("Desktop notes and care plans", "Authorised staff can use the SecureObs staff website for longer notes and care-plan writing. Confirm the correct patient and ward before saving."),
        ],
        header=("Module", "How nurses use it"),
    )
    add_screenshot(doc, ASSETS / "screenshots2" / "Enhanced_Obs_Screen.jpg", "Enhanced observations and TESO: higher-support observation workflow.", width=4.5)

    doc.add_page_break()
    doc.add_heading("5. Safety, escalation and patient tasks", level=1)
    add_step_blocks(
        doc,
        [
            ("Record safety incidents promptly", "Use the Safety and escalation centre for incidents, injuries, safeguarding concerns or significant risks."),
            ("Choose severity", "Use RAG severity to make immediate risk visible: red for immediate/serious risk, amber for review, green for monitor."),
            ("Use body map if relevant", "Select affected body areas for injury or physical concern. Add facts and immediate action taken."),
            ("Assign ownership", "Allocate follow-up to a staff member where required."),
            ("Create patient tasks", "Use tasks for clear follow-up actions such as reviews, checks, contact or clinical actions."),
            ("Review and close", "Acknowledge, resolve or complete items only when the action has genuinely been done."),
        ],
    )
    add_callout(
        doc,
        "Escalation reminder",
        "SecureObs records and highlights risk. It does not replace emergency response, safeguarding policy, nurse-in-charge escalation or local clinical judgement.",
        fill=WARN,
    )

    doc.add_heading("6. End-of-shift handover", level=1)
    add_step_blocks(
        doc,
        [
            ("Open Shift handover", "The handover screen generates evidence from recorded observations, food/fluid, medication, incidents and tasks."),
            ("Read each patient summary", "Check movement, presentation, food/fluid, NEWS2, medication, incidents and outstanding tasks."),
            ("Add staff handover notes", "Add clinical context for the incoming shift. Keep it factual, relevant and concise."),
            ("Review ward summary", "Confirm the overall summary reflects the shift and any unresolved risks or tasks."),
            ("Print draft if needed", "Use Print draft for review or local paper processes."),
            ("Review and sign", "Only sign when the handover is accurate. Signed handovers are retained as clinical records and can be printed/shared as PDF."),
        ],
    )
    add_screenshot(doc, ASSETS / "screenshots2" / "Shift_Handover.jpg", "Shift handover: evidence-based patient summaries with staff handover notes.", width=4.65)

    doc.add_heading("7. Sync, session lock and safe working", level=1)
    add_two_col_table(
        doc,
        [
            ("Sync issues", "Open the sync status to see which record is waiting and the human-readable failure reason. Do not assume a queued record has reached the server."),
            ("Offline use", "Supported records can queue locally. When connectivity returns, sign in with an authorised staff session, retry and confirm SecureObs reports successful synchronisation."),
            ("Session lock", "If the tablet is unused, it locks after the ward timeout. Activity resets the countdown."),
            ("Temporary session ending", "Five minutes before a temporary allocation expires, save notes and unfinished work. SecureObs signs the worker out at the saved end time."),
            ("Wrong patient selected", "Stop before saving, reselect the correct patient and confirm room/hospital number."),
            ("Shared tablet", "Sign out or let the session lock before handing the tablet to another staff member."),
            ("Clinical disagreement", "Record facts and escalate to nurse in charge rather than editing around a concern."),
        ],
        header=("Area", "Nurse action"),
    )

    doc.add_heading("8. Nurse shift checklist", level=1)
    add_bullets(
        doc,
        [
            "Sign in and confirm the correct ward.",
            "Check ward overview for overdue checks, active incidents, tasks and sync issues.",
            "Record each observation against the correct patient.",
            "Use missed/overdue reasons when required.",
            "Where verification is enabled, scan the correct NFC/QR tag and confirm direct visual observation.",
            "Complete NEWS2, food/fluid, enhanced observations or medication records where needed.",
            "Record safety incidents and patient tasks promptly.",
            "Add patient notes or care-plan updates when clinical context is needed.",
            "Review and sign shift handover before leaving the shift.",
            "Check there are no unresolved sync issues before closing the app.",
        ]
    )

    doc.add_page_break()
    doc.add_heading("9. Troubleshooting", level=1)
    add_two_col_table(
        doc,
        [
            ("Patient not visible", "Check ward selection and patient list. Ask manager/admin to check patient setup if still missing."),
            ("Feature missing", "The ward module may be switched off. Ask a manager to check Ward settings."),
            ("Observation saved but sync issue appears", "Open sync status, retry when online, or escalate if it needs review."),
            ("Upload requests sign-in", "Sign in again with an authorised NFC card or STAFFCODE/PIN, retry the named item and confirm that it leaves the queue."),
            ("Cannot access medication", "Check role permissions and whether Medication chart is enabled for the ward."),
            ("Handover looks incomplete", "Check that observations and related records were saved during the shift; add staff notes for context."),
            ("Locked out", "Sign in again with NFC or STAFFCODE/PIN. If locked after failed attempts, ask nurse in charge/manager to unlock access."),
        ],
        header=("Issue", "What to do"),
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
