from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "website" / "training" / "SecureObs_Manager_Training_Manual.docx"
ASSETS = ROOT / "website" / "assets"

NAVY = RGBColor(8, 42, 88)
BLUE = RGBColor(4, 93, 185)
TEAL = RGBColor(8, 167, 181)
MUTED = RGBColor(72, 98, 116)
LIGHT_BLUE = "EAF6FB"
LIGHT_TEAL = "E8F8FA"
LIGHT_GREY = "F3F8FA"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D6E5EC", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_run(run, *, size=None, color=None, bold=None, italic=None, font="Arial") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(18, 36, 46)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 16, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11, NAVY, 8, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("SecureObs Manager Training Manual")
    set_run(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SecureObs • Manager route, staff setup and ward settings")
    set_run(run, size=8, color=MUTED, bold=True)


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    table.autofit = False
    table.columns[0].width = Inches(3.9)
    table.columns[1].width = Inches(2.9)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, "FFFFFF", "0")
            set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    logo = ASSETS / "secureobs-logo-new.png"
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(10)
    p.add_run().add_picture(str(logo), width=Inches(1.05))

    kicker = left.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    r = kicker.add_run("MANAGER TRAINING GUIDE")
    set_run(r, size=10, color=TEAL, bold=True)

    title = left.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 0.95
    r = title.add_run("Setting up staff, access and ward settings")
    set_run(r, size=29, color=NAVY, bold=True)

    lead = left.add_paragraph()
    lead.paragraph_format.space_after = Pt(15)
    r = lead.add_run(
        "A practical manager route manual for signing in, creating permanent staff records, writing NFC staff tags and configuring ward modules safely."
    )
    set_run(r, size=11, color=MUTED, bold=True)

    chips = [
        "Manager/admin access",
        "NFC and PIN sign-in",
        "Ward module switches",
        "Session timeout",
        "Staff rota visibility",
    ]
    chip_table = left.add_table(rows=0, cols=1)
    set_table_width(chip_table, 5200)
    for chip in chips:
        row = chip_table.add_row()
        cell = row.cells[0]
        set_cell_shading(cell, LIGHT_TEAL)
        set_cell_border(cell, "BDECF2")
        set_cell_margins(cell, top=80, bottom=80, start=150, end=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(chip)
        set_run(r, size=9.5, color=NAVY, bold=True)

    image = ASSETS / "screenshots2" / "Management_Settings.jpg"
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Inches(2.75))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(8)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Version 1.0 • July 2026 • For SecureObs managers, ward managers and system administrators")
    set_run(r, size=9, color=MUTED, bold=True)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "C9E5EF")
    set_cell_margins(cell, top=150, bottom=150, start=170, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run(r, size=10.5, color=NAVY, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run(r, size=9.5, color=RGBColor(38, 72, 89), bold=False)


def add_two_col_table(doc: Document, rows: list[tuple[str, str]], widths=(2.35, 4.1), header=None) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    for col, width in zip(table.columns, widths):
        col.width = Inches(width)
    if header:
        row = table.add_row()
        for i, text in enumerate(header):
            cell = row.cells[i]
            set_cell_shading(cell, "0A2F63")
            set_cell_border(cell, "0A2F63")
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run(r, size=9, color=WHITE, bold=True)
    for label, detail in rows:
        row = table.add_row()
        for i, text in enumerate((label, detail)):
            cell = row.cells[i]
            set_cell_shading(cell, "FFFFFF" if i else LIGHT_GREY)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run(r, size=9, color=NAVY if i == 0 else RGBColor(31, 54, 66), bold=i == 0)


def add_step_blocks(doc: Document, steps: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    table.columns[0].width = Inches(0.65)
    table.columns[1].width = Inches(5.8)
    for idx, (title, body) in enumerate(steps, 1):
        row = table.add_row()
        number_cell, body_cell = row.cells
        set_cell_shading(number_cell, "0A6C85")
        set_cell_border(number_cell, "0A6C85")
        set_cell_margins(number_cell, top=140, bottom=140, start=100, end=100)
        p = number_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(idx))
        set_run(r, size=14, color=WHITE, bold=True)

        set_cell_shading(body_cell, "FFFFFF")
        set_cell_border(body_cell)
        set_cell_margins(body_cell, top=120, bottom=120, start=150, end=150)
        p = body_cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        set_run(r, size=10, color=NAVY, bold=True)
        p = body_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(body)
        set_run(r, size=9, color=RGBColor(45, 70, 82))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run(r, size=9.7, color=RGBColor(30, 55, 68))


def add_screenshot(doc: Document, path: Path, caption: str, width=6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run(r, size=8.5, color=MUTED, italic=True)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def build() -> None:
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_page_break(doc)

    doc.add_heading("1. What this manager manual covers", level=1)
    add_callout(
        doc,
        "Manager route purpose",
        "The manager route is used to control access, staff records, ward module switches and the operational settings that shape what staff can see and do on a ward tablet.",
    )
    add_two_col_table(
        doc,
        [
            ("Who should use this guide", "Ward managers, managers with admin rights, super admins and implementation leads."),
            ("Before starting", "Confirm the site and ward exist, agree local STAFFCODE format, and prepare any blank NFC staff tags."),
            ("Important safety point", "Create staff with the minimum access they need. Inactivate records promptly when a staff member leaves or no longer needs access."),
        ],
        header=("Item", "Manager guidance"),
    )

    doc.add_heading("2. Manager sign-in route", level=1)
    add_step_blocks(
        doc,
        [
            ("Open SecureObs", "The first screen is the staff access area. The current session strip shows selected staff, site and ward."),
            ("Sign in with NFC or STAFFCODE/PIN", "Use Scan card for a written NFC staff tag, or enter STAFFCODE and PIN. Permanent staff can use the PIN route when NFC is not available."),
            ("Change PIN if prompted", "If the app says PIN change required, enter the current PIN, then the new 4–6 digit PIN twice before continuing."),
            ("Choose site and ward", "Managers with wider access can select the correct site and ward before opening settings or starting the ward workflow."),
            ("Open settings", "Use SecureObs admin for organisation-level setup, or Ward settings for staff access, modules and ward-specific controls."),
        ],
    )
    add_screenshot(doc, ASSETS / "screenshots2" / "sign_In_Screen.jpg", "Staff access screen: NFC, STAFFCODE/PIN and current session selection.", width=4.4)

    doc.add_heading("3. SecureObs admin screen", level=1)
    add_two_col_table(
        doc,
        [
            ("Audit log", "Search staff lookups, observation saves, medication activity, settings changes and failed access."),
            ("Company branding", "Upload or remove the organisation logo shown in the app header."),
            ("NFC staff card format", "Controls how staff cards are read. Keep {STAFFCODE} in the format, e.g. passcode={STAFFCODE}."),
            ("Sites", "Create site or care home records. Select a site before adding wards beneath it."),
            ("Wards", "Create ward records, assign service type, set the initial observation interval and optionally create a manager account at the same time."),
            ("Service type", "Helps identify the setting type and supports local configuration decisions."),
            ("Observation interval", "Sets the starting intermittent observation frequency for new wards."),
        ],
        header=("Admin section", "What it does"),
    )
    add_callout(
        doc,
        "NFC format warning",
        "Only change the NFC staff card format if you understand the tag data being written. Existing fallback formats and plain STAFFCODE still work, but a mismatched format can stop scans from resolving cleanly.",
        fill="FFF4D6",
    )

    doc.add_heading("4. Adding or updating staff", level=1)
    add_step_blocks(
        doc,
        [
            ("Search before adding", "Search site staff by name, STAFFCODE or role. If the person already exists, select them and update ward access instead of creating a duplicate."),
            ("Choose Add new", "Use Add new to clear the draft and create a permanent staff record."),
            ("Enter identity details", "Complete Staff name, STAFFCODE, Designation and optional initial PIN. If no initial PIN is entered, the default is 1111."),
            ("Select role", "Choose nurse, HCF, OT, security or doctor. The role controls which features the staff member can use."),
            ("Set ward access", "Select at least one ward. Multi-ward staff can be given access to more than one ward at the same site or organisation."),
            ("Set status", "Active staff can sign in. Inactive staff should not be able to continue working from that record."),
            ("Doctor prescribing", "Only doctor records can be marked Can prescribe. Leave this off unless prescribing access is intended."),
            ("Save the record", "Tap Add staff member or Update staff member. A confirmation appears once the save has been accepted locally and queued/synced as needed."),
        ],
    )
    add_two_col_table(
        doc,
        [
            ("nurse", "Clinical nursing access. Can use core ward workflows and manager-permitted clinical screens."),
            ("hcf", "Healthcare/support worker access. Cannot be Nurse in Charge or Medication Nurse in rota assignment."),
            ("ot", "Occupational therapy role where enabled by the service."),
            ("security", "Security-focused access for security checks and related ward activity."),
            ("doctor", "Medical role. Can be granted prescribing access with the Can prescribe control."),
            ("manager / super admin", "Higher access roles used for settings, setup and governance controls."),
        ],
        header=("Role", "Typical use"),
    )

    doc.add_heading("5. Writing NFC staff tags", level=1)
    add_callout(
        doc,
        "Permanent staff only",
        "NFC staff tags are intended for permanent staff records. Save the staff member first, then use Write NFC tag from the confirmation prompt or the NFC staff tag panel.",
    )
    add_step_blocks(
        doc,
        [
            ("Save staff first", "Enter staff details and save the staff member. The app needs the staff record before writing the tag."),
            ("Tap Write NFC tag", "The panel shows the payload that will be written, using the organisation NFC format."),
            ("Hold the blank tag to the device", "Keep the tag still until the app confirms it has been written."),
            ("Test the tag", "Sign out and scan the card on the staff access screen to confirm it selects the correct staff member."),
        ],
    )

    doc.add_heading("6. Ward settings map", level=1)
    add_two_col_table(
        doc,
        [
            ("Landing page after login", "Choose whether staff first see Ward overview or Patient checks after selecting the ward."),
            ("NEWS2", "Shows NEWS2 charting and score history."),
            ("Enhanced observations", "Enables TESO start, plan of care and enhanced observation recording."),
            ("Security checks", "Enables ward checkpoint recording and security check workflows."),
            ("Configure security checks", "Opens the setup screen for the ward’s security check areas/items."),
            ("Medication chart", "Enables medication due prompts, prescribing and administration recording where roles permit."),
            ("Assessment forms", "Enables care-home style forms, signatures and printable assessments."),
            ("Food and fluid chart", "Enables meal, snack and drink intake monitoring."),
            ("Intermittent observation interval", "Sets the default routine observation frequency. Quick buttons are 15, 30 and 60 minutes, with -5/+5 controls."),
            ("Staff session timeout", "Controls the lock countdown after inactivity. The countdown starts after 2 minutes without touch or typing."),
            ("Staff rota", "Shows or hides the staff rota page for that ward."),
        ],
        header=("Setting", "Manager explanation"),
    )

    add_page_break(doc)
    doc.add_heading("7. Recommended manager workflow", level=1)
    add_bullets(
        doc,
        [
            "Start with organisation setup: add the logo, confirm NFC format, create sites and create wards.",
            "Create or update manager accounts before wider staff onboarding.",
            "Add permanent staff records, assign ward access and set active status.",
            "Write and test NFC tags while the staff member is present.",
            "Turn on only the modules the ward is ready to use.",
            "Set observation interval and session timeout according to local policy.",
            "Enable staff rota only once the ward is ready to allocate roles inside SecureObs.",
            "Use the audit log when checking sign-ins, failed access or settings changes.",
        ]
    )

    doc.add_heading("8. Safety checks before go-live", level=1)
    checklist = [
        ("Sites and wards", "Every live ward exists under the correct site."),
        ("Manager access", "At least two appropriate managers or admins can access settings."),
        ("Staff access", "Permanent staff have correct roles, ward access and active status."),
        ("NFC tags", "Tags are written, tested and labelled according to local process."),
        ("Session timeout", "Timeout is agreed locally. Minimum effective lock is 2 minutes grace plus the selected timeout."),
        ("Modules", "Only approved modules are enabled for the ward."),
        ("Rota", "Rota is enabled only if staff are expected to use it."),
        ("Audit", "Managers know where to view audit events and failed access."),
    ]
    add_two_col_table(doc, checklist, header=("Check", "What good looks like"))

    doc.add_heading("9. Troubleshooting", level=1)
    add_two_col_table(
        doc,
        [
            ("Staff cannot sign in", "Check active status, STAFFCODE spelling, PIN, ward access and whether the staff member is locked after failed attempts."),
            ("NFC tag reads the wrong code", "Check the organisation NFC staff card format and rewrite the tag after confirming STAFFCODE."),
            ("Manager cannot edit settings", "Confirm the signed-in staff member has manager or admin access and has completed any required PIN change."),
            ("Rota save does not appear", "Check the staff member has manager/nurse/super admin access and that the tablet sync queue is clear."),
            ("Feature is missing", "Open Ward settings and confirm the module is switched on for that ward."),
            ("App locks after inactivity", "This is expected. Any touch or typing resets the countdown. The minimum effective timeout is currently 17 minutes."),
        ],
        header=("Issue", "Manager action"),
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
